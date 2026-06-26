import re
from docx import Document
import json
from typing import Dict, Any


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a .docx file bytes object."""
    import io
    doc = Document(io.BytesIO(file_bytes))
    
    full_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Also extract from tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in full_text:
                    full_text.append(cell_text)
    
    return "\n".join(full_text)


def parse_resume_with_llm(raw_text: str, token: str = None) -> Dict[str, Any]:
    """Use the configured LLM to parse raw resume text into structured JSON."""
    from app.tailor import generate_llm_content, clean_json_response

    prompt = f"""
You are an expert resume parser. Below is the raw text extracted from a candidate's resume Word document.
Please analyze the text carefully and extract structured resume information.

Resume Text:
---
{raw_text[:8000]}
---

Return a JSON object with EXACTLY this structure:
{{
    "summary": "Professional summary or objective statement. If not found, write a short 2-3 sentence professional summary based on the experience.",
    "experience": [
        {{
            "company": "Company name",
            "title": "Job title",
            "dates": "Start - End dates (e.g. Jan 2023 - Present)",
            "location": "City, Country or Remote",
            "bullets": [
                "Achievement or responsibility bullet point 1",
                "Achievement or responsibility bullet point 2"
            ]
        }}
    ],
    "education": [
        {{
            "school": "University or institution name",
            "degree": "Degree name and field of study",
            "dates": "Year or date range"
        }}
    ],
    "skills": ["Skill 1", "Skill 2", "Skill 3"]
}}

Rules:
- Extract ALL work experiences, internships, and roles.
- Extract ALL education entries including certifications and professional qualifications (CFA, CPA, etc.).
- Skills should be a flat array of individual skill strings.
- If a section is not found, return an empty array [] for it.
- Return ONLY the raw JSON object — no markdown code fences, no explanation text.
"""

    response_text = generate_llm_content(prompt, token=token).strip()
    response_text = clean_json_response(response_text)
    return json.loads(response_text)


def parse_resume_heuristic(raw_text: str) -> Dict[str, Any]:
    """
    Fallback heuristic parser when Gemini API key is not set.
    Extracts skills and provides a basic structure.
    """
    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]

    # Common section headers to detect
    section_keywords = {
        "summary": ["summary", "profile", "objective", "about me", "professional summary"],
        "experience": ["experience", "work history", "employment", "positions", "roles"],
        "education": ["education", "qualifications", "academic", "degree", "university"],
        "skills": ["skills", "technical skills", "competencies", "tools", "technologies"],
    }

    # Build skills list from comma-separated lines
    skills = []
    for line in lines:
        if "," in line and len(line) < 200:
            candidates = [s.strip() for s in line.split(",")]
            if all(len(c) < 40 for c in candidates):
                skills.extend(candidates)

    return {
        "summary": lines[0] if lines else "",
        "experience": [],
        "education": [],
        "skills": list(set([s for s in skills if s]))[:20],
        "_raw_text": raw_text,
        "_parse_warning": "Gemini API key not set. Resume was extracted as raw text. Please configure your API key and re-upload for intelligent parsing."
    }
