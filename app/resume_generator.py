import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume_docx(resume_data: dict, filepath: str) -> str:
    """
    Generates a clean, professional Word document resume from structured resume JSON.
    """
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Setup document base font to Calibri/Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # 1. Candidate Name (Header)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name_run = p_name.add_run(resume_data.get("name", "DHAVAL MANAWAT").upper())
    p_name_run.bold = True
    p_name_run.font.size = Pt(18)
    
    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = [
        resume_data.get("email", "dhava1.m@example.com"),
        resume_data.get("phone", "+91 98765 43210"),
        "Mumbai, India"
    ]
    p_contact.add_run(" | ".join(contact_parts))
    p_contact.paragraph_format.space_after = Pt(12)
    
    # Helper to add section headers with divider lines
    def add_section_header(title):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        
        # Add thin light-gray divider border bottom
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 6/8 pt line thickness
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')  # Light gray
        pBdr.append(bottom)
        pPr.append(pBdr)
        
    # 2. Professional Summary
    summary = resume_data.get("summary", "")
    if summary:
        add_section_header("Professional Summary")
        doc.add_paragraph(summary)
        
    # 3. Experience
    experience = resume_data.get("experience", [])
    if experience:
        add_section_header("Work Experience")
        for exp in experience:
            p_job = doc.add_paragraph()
            p_job.paragraph_format.space_before = Pt(6)
            p_job.paragraph_format.space_after = Pt(2)
            p_job.paragraph_format.keep_with_next = True
            
            # Title & Company
            title_run = p_job.add_run(exp.get("title", ""))
            title_run.bold = True
            p_job.add_run(" at ")
            company_run = p_job.add_run(exp.get("company", ""))
            company_run.italic = True
            
            # Dates & Location (Right-aligned tab space or just separate line)
            # Subtitle line
            p_meta = doc.add_paragraph()
            p_meta.paragraph_format.space_after = Pt(2)
            p_meta.paragraph_format.keep_with_next = True
            meta_run = p_meta.add_run(f"{exp.get('dates', '')} | {exp.get('location', '')}")
            meta_run.font.size = Pt(9.5)
            meta_run.font.color.rgb = None # default dark gray
            
            # Bullet points
            bullets = exp.get("bullets", [])
            for bullet in bullets:
                # Clean bullets from leading bullet characters
                cleaned_bullet = bullet.strip().lstrip("▸\t-•*").strip()
                if cleaned_bullet:
                    doc.add_paragraph(cleaned_bullet, style='List Bullet')
                    
    # 4. Education
    education = resume_data.get("education", [])
    if education:
        add_section_header("Education")
        for edu in education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_before = Pt(4)
            p_edu.paragraph_format.space_after = Pt(2)
            
            degree_run = p_edu.add_run(edu.get("degree", ""))
            degree_run.bold = True
            p_edu.add_run(" | ")
            school_run = p_edu.add_run(edu.get("school", ""))
            
            if edu.get("dates"):
                p_edu.add_run(f" ({edu.get('dates')})")
                
    # 5. Skills
    skills = resume_data.get("skills", [])
    if skills:
        add_section_header("Skills & Certifications")
        p_skills = doc.add_paragraph()
        p_skills.paragraph_format.space_before = Pt(4)
        skills_str = ", ".join(skills)
        p_skills.add_run(skills_str)
        
    # Save the document
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    return filepath
